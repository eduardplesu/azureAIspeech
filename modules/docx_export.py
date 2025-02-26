# modules/docx_export.py
from docx import Document

def ticks_to_time(ticks):
    """
    Converts ticks (100-nanosecond intervals) to a formatted time string HH:MM:SS.mmm.
    """
    seconds = ticks / 10_000_000  # convert ticks to seconds
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:06.3f}"

def export_transcription_to_docx(transcription_results, analysis_text=None, output_filename="transcription.docx", cleaned_transcription=None):
    """
    Exports transcription results to a DOCX file.
    Each transcription segment includes speaker information, text, and time details.
    If 'speaker_name' exists, it is used; otherwise, 'speaker_id' is shown.
    If cleaned_transcription is provided, it is added as a separate section.
    If analysis_text is provided, it is added as an Analysis section.
    """
    document = Document()
    document.add_heading("Transcription", level=0)
    
    for result in transcription_results:
        speaker = result.get("speaker_name", result.get("speaker_id", "Unknown"))
        text = result.get("text", "")
        offset = result.get("offset", 0)
        duration = result.get("duration", 0)
        start_time = ticks_to_time(offset)
        duration_time = ticks_to_time(duration)
        document.add_paragraph(f"Speaker {speaker}: {text}")
        document.add_paragraph(f"(Start Time: {start_time}, Duration: {duration_time})", style="Intense Quote")
    
    if cleaned_transcription:
        document.add_page_break()
        document.add_heading("Cleaned Transcription", level=0)
        document.add_paragraph(cleaned_transcription)
    
    if analysis_text:
        document.add_page_break()
        document.add_heading("Analysis", level=0)
        document.add_paragraph(analysis_text)
    
    document.save(output_filename)
    return output_filename
